"""
AI-Powered Maintenance Query Agent — Streamlit App
File: streamlit_rag_maintenance_app.py

Features:
- Upload equipment manuals (PDF, DOCX, TXT, CSV) and logs
- Chunk documents and store embeddings (OpenAI or local sentence-transformers)
- Vector store using FAISS
- Retrieval-Augmented Generation (RAG) using OpenAI Responses API
- Simple Streamlit UI: upload, index, query, show retrieved chunks and answer

Requirements (put into requirements.txt):
streamlit
openai
faiss-cpu
langchain
python-docx
pdfplumber
sentence-transformers
tqdm
pandas
transformers
torch
duckdb  # optional for storing metadata

NOTE: Set environment variable OPENAI_API_KEY when using OpenAI embeddings / LLM.

Usage:
streamlit run streamlit_rag_maintenance_app.py

"""

import os
import tempfile
import uuid
import io
from typing import List, Dict, Any

import streamlit as st
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.docstore.document import Document
from langchain.embeddings import OpenAIEmbeddings
from langchain.vectorstores import FAISS

# Optional local embedding backend
try:
    from sentence_transformers import SentenceTransformer
    LOCAL_SBM_AVAILABLE = True
except Exception:
    LOCAL_SBM_AVAILABLE = False

# File parsing
import pdfplumber
import docx
import pandas as pd
from tqdm.auto import tqdm

# OpenAI client
try:
    from openai import OpenAI
    OPENAI_CLIENT_AVAILABLE = True
except Exception:
    OPENAI_CLIENT_AVAILABLE = False

# ---------------------------
# Helpers: File loaders
# ---------------------------

def load_pdf(file_bytes: bytes) -> str:
    text_pages = []
    with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
        for p in pdf.pages:
            try:
                text = p.extract_text()
            except Exception:
                text = None
            if text:
                text_pages.append(text)
    return "\n\n".join(text_pages)


def load_docx(file_bytes: bytes) -> str:
    doc = docx.Document(io.BytesIO(file_bytes))
    paragraphs = [p.text for p in doc.paragraphs if p.text and p.text.strip()]
    return "\n\n".join(paragraphs)


def load_txt(file_bytes: bytes) -> str:
    return file_bytes.decode("utf-8", errors="ignore")


def load_csv(file_bytes: bytes) -> str:
    df = pd.read_csv(io.BytesIO(file_bytes))
    # Convert rows to human readable text
    rows = []
    for _, r in df.iterrows():
        rows.append("; ".join([f"{c}: {r[c]}" for c in df.columns]))
    return "\n\n".join(rows)


def extract_text_from_upload(uploaded_file) -> str:
    """Detect type and extract text from uploaded file-like (streamlit UploadedFile)."""
    if not uploaded_file:
        return ""
    name = uploaded_file.name.lower()
    data = uploaded_file.read()
    if name.endswith(".pdf"):
        return load_pdf(data)
    elif name.endswith(".docx"):
        return load_docx(data)
    elif name.endswith(".txt"):
        return load_txt(data)
    elif name.endswith(".csv"):
        return load_csv(data)
    else:
        # Fallback try decode
        try:
            return data.decode("utf-8", errors="ignore")
        except Exception:
            return ""


# ---------------------------
# Embedding backends
# ---------------------------

class LocalSentenceTransformerEmbedder:
    def __init__(self, model_name: str = "all-MiniLM-L6-v2"):
        if not LOCAL_SBM_AVAILABLE:
            raise RuntimeError("sentence-transformers not available")
        self.model = SentenceTransformer(model_name)

    def embed_documents(self, texts: List[str]) -> List[List[float]]:
        return self.model.encode(texts, show_progress_bar=False)

    def embed_query(self, text: str) -> List[float]:
        return self.model.encode([text])[0]


class OpenAIEmbedder:
    def __init__(self, model_name: str = "text-embedding-3-large"):
        if not OPENAI_CLIENT_AVAILABLE:
            raise RuntimeError("openai client not available")
        # using langchain OpenAIEmbeddings for convenience
        self.embedder = OpenAIEmbeddings(model=model_name)

    def embed_documents(self, texts: List[str]) -> List[List[float]]:
        return self.embedder.embed_documents(texts)

    def embed_query(self, text: str) -> List[float]:
        return self.embedder.embed_query(text)


# ---------------------------
# Chunking utilities
# ---------------------------

def chunk_texts(texts: List[str], chunk_size: int = 800, chunk_overlap: int = 150) -> List[Document]:
    splitter = RecursiveCharacterTextSplitter(chunk_size=chunk_size, chunk_overlap=chunk_overlap)
    docs: List[Document] = []
    for i, t in enumerate(texts):
        if not t or not t.strip():
            continue
        pieces = splitter.split_text(t)
        for j, p in enumerate(pieces):
            metadata = {"source": f"uploaded_{i}", "chunk": j}
            docs.append(Document(page_content=p, metadata=metadata))
    return docs


# ---------------------------
# RAG: retrieval + LLM answer
# ---------------------------

def build_prompt(question: str, retrieved_docs: List[Document], max_context_chars: int = 4000) -> str:
    # Prepare a compact context: include source metadata and content truncated
    ctx_parts = []
    total_chars = 0
    for d in retrieved_docs:
        snippet = d.page_content
        # truncate to avoid overly long prompt
        if len(snippet) > 1000:
            snippet = snippet[:1000] + "..."
        entry = f"Source: {d.metadata.get('source', 'unknown')} | chunk: {d.metadata.get('chunk', '?')}\n{snippet}"
        total_chars += len(entry)
        if total_chars > max_context_chars:
            break
        ctx_parts.append(entry)

    context = "\n\n---\n\n".join(ctx_parts) if ctx_parts else ""

    prompt = f"You are an experienced industrial maintenance engineer. Use ONLY the context provided to answer the question. If the exact answer is not in the context, say 'Not found in uploaded manuals/logs; provide recommended general troubleshooting steps.'\n\nCONTEXT:\n{context}\n\nQUESTION:\n{question}\n\nAnswer with concise, step-by-step troubleshooting instructions and suggested safety precautions."
    return prompt


def query_openai_responses(prompt: str, model: str = "gpt-5.1") -> str:
    if not OPENAI_CLIENT_AVAILABLE:
        raise RuntimeError("OpenAI client not available in environment.")
    client = OpenAI()
    # Use the Responses API
    resp = client.responses.create(model=model, input=prompt)
    # The shape may vary; try common access patterns
    if hasattr(resp, "output"):
        # lang-like wrapper
        out = resp.output
        if isinstance(out, list):
            # join content pieces
            texts = []
            for item in out:
                if isinstance(item, dict) and "content" in item:
                    # older formats
                    texts.append(item.get("content", ""))
                elif hasattr(item, "text"):
                    texts.append(item.text)
            return "\n\n".join(texts).strip()
        elif isinstance(out, str):
            return out

    # Fallback
    try:
        return resp.output_text
    except Exception:
        return str(resp)


# ---------------------------
# App: Streamlit UI
# ---------------------------

st.set_page_config(page_title="Maintenance RAG Agent", layout="wide")
st.title("🔧 AI-Powered Maintenance Query Agent — Streamlit")

# Sidebar: settings
st.sidebar.header("Index / Embedding Settings")
embedding_backend = st.sidebar.selectbox("Embedding backend", options=["OpenAI (cloud)", "Local (sentence-transformers)"], index=0 if OPENAI_CLIENT_AVAILABLE else 1)
chunk_size = st.sidebar.slider("Chunk size (chars)", min_value=200, max_value=2000, value=800, step=100)
chunk_overlap = st.sidebar.slider("Chunk overlap (chars)", min_value=0, max_value=500, value=150, step=50)
use_faiss_memory = True

# Upload area
st.markdown("### 1) Upload manuals & logs")
uploaded_files = st.file_uploader("Upload PDF / DOCX / TXT / CSV files", accept_multiple_files=True)

if 'vector_store' not in st.session_state:
    st.session_state['vector_store'] = None
    st.session_state['documents'] = []
    st.session_state['index_meta'] = {}

col1, col2 = st.columns([2, 1])
with col1:
    if uploaded_files:
        st.write(f"{len(uploaded_files)} files uploaded:")
        for f in uploaded_files:
            st.write(f"- {f.name} ({f.type})")

with col2:
    if st.button("Index uploaded files"):
        # Extract texts
        with st.spinner("Extracting text from files..."):
            text_blobs = []
            filenames = []
            for f in tqdm(uploaded_files, desc="parsing"):
                txt = extract_text_from_upload(f)
                if txt and txt.strip():
                    text_blobs.append(txt)
                    filenames.append(f.name)

        if not text_blobs:
            st.warning("No readable text extracted from uploaded files. Try different files or formats.")
        else:
            # Chunk
            with st.spinner("Chunking documents..."):
                docs = chunk_texts(text_blobs, chunk_size=chunk_size, chunk_overlap=chunk_overlap)
            st.session_state['documents'] = docs
            st.session_state['index_meta'] = {"filenames": filenames}

            # Build embeddings & FAISS
            with st.spinner("Creating embeddings and FAISS index (this may take a while)..."):
                texts = [d.page_content for d in docs]
                if embedding_backend.startswith("OpenAI"):
                    if not OPENAI_CLIENT_AVAILABLE:
                        st.error("OpenAI client library not installed. Please install 'openai' and ensure OPENAI_API_KEY is set.")
                    embedder = OpenAIEmbedder()
                    # Use langchain's FAISS builder for convenience
                    vector_store = FAISS.from_texts(texts, embedding=OpenAIEmbeddings(model="text-embedding-3-large"))
                else:
                    if not LOCAL_SBM_AVAILABLE:
                        st.error("Local sentence-transformers not available. Install 'sentence-transformers'.")
                        st.stop()
                    model_name = st.sidebar.text_input("Local SBERT model", value="all-MiniLM-L6-v2")
                    local_embed = LocalSentenceTransformerEmbedder(model_name=model_name)
                    vectors = local_embed.embed_documents(texts)
                    vector_store = FAISS.from_texts(texts, embedding=local_embed)  # langchain can accept custom embedder-like object

                st.session_state['vector_store'] = vector_store
                st.success(f"Indexed {len(docs)} chunks from {len(text_blobs)} documents.")

# Query area
st.markdown("### 2) Ask natural language questions")
query = st.text_area("Enter your maintenance question", height=120)
k = st.slider("Number of chunks to retrieve (k)", min_value=1, max_value=10, value=5)
selected_model = st.selectbox("LLM for answers", options=["gpt-5.1", "gpt-4o"], index=0)

if st.button("Ask"):
    if not query or not query.strip():
        st.warning("Please type a question.")
    elif not st.session_state.get('vector_store'):
        st.warning("No index found. Upload and index files first.")
    else:
        with st.spinner("Retrieving relevant context from FAISS..."):
            vs: FAISS = st.session_state['vector_store']
            retrieved = vs.similarity_search(query, k=k)

        # Show retrieved snippets in expandable
        st.markdown("#### Retrieved context")
        for i, r in enumerate(retrieved):
            with st.expander(f"Chunk {i+1} — source: {r.metadata.get('source', 'unknown')}"):
                st.write(r.page_content[:2000])

        # Build prompt and query LLM
        prompt = build_prompt(query, retrieved, max_context_chars=6000)
        st.markdown("#### Generated Answer")
        try:
            answer = query_openai_responses(prompt, model=selected_model)
            st.text_area("Answer", value=answer, height=300)
        except Exception as e:
            st.error(f"Error querying LLM: {e}")

# Sidebar: session actions
st.sidebar.header("Session")
if st.sidebar.button("Clear index"):
    st.session_state['vector_store'] = None
    st.session_state['documents'] = []
    st.session_state['index_meta'] = {}
    st.success("Cleared index from session")

st.sidebar.markdown("---")
st.sidebar.write("Tips:")
st.sidebar.write("- Upload the latest equipment manuals and logs (PDF/CSV).")
st.sidebar.write("- Use OpenAI embeddings for higher-quality semantic search, or local SBERT for offline.")
st.sidebar.write("- Record actions suggested by the assistant into your CMMS (Computerized Maintenance Management System).")

# Footer
st.markdown("---")
st.caption("This app is a starter template. For production use, add authentication, logging, safe-action checks, and audit trails.")

# End of file
